import streamlit as st
from google.generativeai import configure
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import PromptTemplate
from docx import Document
from io import BytesIO
import base64
import os
from dotenv import load_dotenv

st.set_page_config(layout="wide")

load_dotenv()
configure(api_key=os.getenv("GOOGLE_API_KEY"))

st.sidebar.image("medicebot.webp", width=300)  

st.title("MediceBot - AI Agents to Empower Doctors 💡")
st.write("Welcome to the MediceBot app, powered by AI. Please provide the information for diagnosis and treatment recommendations.")

# Input fields for user details
gender = st.sidebar.selectbox('Select Gender', ('Male', 'Female', 'Other'))
age = st.sidebar.number_input('Enter Age', min_value=0, max_value=120, value=25)
symptoms = st.sidebar.text_area('Enter Symptoms', 'e.g., fever, cough, headache')
medical_history = st.sidebar.text_area('Enter Medical History', 'e.g., diabetes, hypertension')

llm = ChatGoogleGenerativeAI(model="gemini-pro", temperature=0.3)

# Define the prompts for diagnosis and treatment
diagnosis_prompt = PromptTemplate(
    template=(
        "You are a Medical Diagnostician. Analyze the following details:\n"
        "Gender: {gender}\n"
        "Age: {age}\n"
        "Symptoms: {symptoms}\n"
        "Medical History: {medical_history}\n\n"
        "Provide a preliminary diagnosis with possible conditions based on this information."
    ),
    input_variables=["gender", "age", "symptoms", "medical_history"]
)

treatment_prompt = PromptTemplate(
    template=(
        "You are a Treatment Advisor. Based on the diagnosis:\n"
        "{diagnosis}\n"
        "and considering the patient's symptoms and medical history:\n"
        "Symptoms: {symptoms}\n"
        "Medical History: {medical_history}\n\n"
        "Recommend an appropriate treatment plan, including medications, lifestyle changes, and follow-up care."
    ),
    input_variables=["diagnosis", "symptoms", "medical_history"]
)

# Function to generate a Word document
def generate_docx(result):
    doc = Document()
    doc.add_heading('Healthcare Diagnosis and Treatment Recommendations', 0)
    doc.add_paragraph(result)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# Function to generate a download link for the Word document
def get_download_link(bio, filename):
    b64 = base64.b64encode(bio.read()).decode()
    return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download Diagnosis and Treatment Plan</a>'

if st.sidebar.button("Get Diagnosis and Treatment Plan"):
    with st.spinner('Generating recommendations...'):
        # Prepare diagnosis input
        diagnosis_input = diagnosis_prompt.format(
            gender=gender,
            age=age,
            symptoms=symptoms,
            medical_history=medical_history
        )
        diagnosis = llm.predict(diagnosis_input)

        # Prepare treatment input
        treatment_input = treatment_prompt.format(
            diagnosis=diagnosis,
            symptoms=symptoms,
            medical_history=medical_history
        )
        treatment_plan = llm.predict(treatment_input)

        # Display results
        result = f"**Diagnosis:**\n{diagnosis}\n\n**Treatment Plan:**\n{treatment_plan}"
        st.write(result)

        # Generate downloadable Word document
        docx_file = generate_docx(result)
        download_link = get_download_link(docx_file, "diagnosis_and_treatment_plan.docx")
        st.markdown(download_link, unsafe_allow_html=True)

# Footer Section
st.markdown(
    """
    <hr>
    <div style="text-align: center;">
        <p>Built with ❤️ by <strong>Kamal Aliyu</strong></p>
        <a href="https://github.com/Alkamal01" target="_blank">
            <img src="https://img.icons8.com/material-outlined/24/github.png" alt="GitHub">
        </a>
        <a href="https://twitter.com/kaftandev" target="_blank">
            <img src="https://img.icons8.com/material-outlined/24/twitter.png" alt="Twitter">
        </a>
        <a href="https://linkedin.com/in/alkamal" target="_blank">
            <img src="https://img.icons8.com/material-outlined/24/linkedin.png" alt="LinkedIn">
        </a>
    </div>
    """,
    unsafe_allow_html=True
)
